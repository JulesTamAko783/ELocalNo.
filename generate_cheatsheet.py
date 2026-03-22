from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# -- Page margins (narrow for 1-page fit) --
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# -- Style helpers --
style = doc.styles['Normal']
style.font.name = 'Consolas'
style.font.size = Pt(8.5)
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(1)

DARK_BG = "1E1E2E"
CODE_BG = "2D2D3D"
ACCENT = "89B4FA"
GREEN = "A6E3A1"
PEACH = "FAB387"
TEXT = "CDD6F4"
MAUVE = "CBA6F7"

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_cell_text(cell, text, bold=False, color=None, size=Pt(8.5), font_name='Consolas'):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        add_cell_text(cell, h, bold=True, color=ACCENT, size=Pt(8))
        set_cell_shading(cell, DARK_BG)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            add_cell_text(cell, val, color=TEXT if c_idx > 0 else GREEN, size=Pt(8))
            set_cell_shading(cell, CODE_BG if r_idx % 2 == 0 else DARK_BG)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12) if level == 0 else Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MAUVE if level == 0 else ACCENT)
    run.font.name = 'Segoe UI'
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    cell.paragraphs[0].clear()
    for i, line in enumerate(code.strip().split('\n')):
        if i > 0:
            cell.paragraphs[0].add_run('\n')
        run = cell.paragraphs[0].add_run(line)
        run.font.size = Pt(7.5)
        run.font.name = 'Consolas'
        run.font.color.rgb = RGBColor.from_string(TEXT)

def add_body(doc, text, color=TEXT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Consolas'
    return p

# ============================================================
# TITLE
# ============================================================
add_heading_styled(doc, "ELOKANO LANGUAGE CHEATSHEET", level=0)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("File extension: .elokano  |  Run: python coderun.py file.elokano --out out.cpp --exe out.exe")
r.font.size = Pt(7.5)
r.font.color.rgb = RGBColor.from_string(PEACH)
r.font.name = 'Consolas'
subtitle.paragraph_format.space_after = Pt(4)

# ============================================================
# DATA TYPES
# ============================================================
add_heading_styled(doc, "DATA TYPES")
styled_table(doc,
    ["Type", "Description", "Default", "Size", "Example"],
    [
        ["Bilang",    "Integer",      "0",     "4 bytes",  'Bilang x dutokan-> 10;'],
        ["Gudua",     "Float/Double", "0.0",   "8 bytes",  'Gudua pi dutokan-> 3.14;'],
        ["Sarsarita", "String",       '""',    "32 bytes", 'Sarsarita s dutokan-> "hi";'],
        ["Pudno",     "Boolean",      "false", "1 byte",   'Pudno f dutokan-> true;'],
    ],
    col_widths=[0.8, 0.8, 0.5, 0.6, 2.6]
)

# ============================================================
# VARIABLES
# ============================================================
add_heading_styled(doc, "VARIABLES")
add_code_block(doc, """\
Bilang age dutokan-> 21;                      // declare + assign
Sarsarita name;                               // declare with default ""
Bilang x dutokan-> 1, y dutokan-> 2, z;       // multi-declare same type
age dutokan-> 25;                             // reassign""")
add_body(doc, 'dutokan-> is the assignment operator. Every statement ends with ;')
add_body(doc, 'Gudua can receive Bilang values; other type mismatches are errors.', PEACH)

# ============================================================
# OPERATORS
# ============================================================
add_heading_styled(doc, "OPERATORS")
styled_table(doc,
    ["Category", "Operators"],
    [
        ["Arithmetic",  "+  -  *  / (float div)  // (int div)  % (mod)"],
        ["Comparison",  "==  !=  >  <  >=  <=   (returns Pudno)"],
        ["String",      "+ (concatenation)"],
    ],
    col_widths=[1.0, 4.3]
)

# ============================================================
# INPUT / OUTPUT
# ============================================================
add_heading_styled(doc, "INPUT / OUTPUT")
add_code_block(doc, """\
Ibaga("Hello World");                          // print with newline
Ibaga("No newline", "");                       // custom line ending
Ibaga("Tab end", "\\t");                        // end with tab
Ibaga();                                       // blank line
Sarsarita name dutokan-> Ikabil("Name: ");     // input with prompt
Sarsarita raw dutokan-> Ikabil;                // input without prompt""")

# ============================================================
# CONTROL FLOW
# ============================================================
add_heading_styled(doc, "CONTROL FLOW  (if / else-if / else)")
styled_table(doc,
    ["Keyword", "Meaning"],
    [
        ["nu",        "if"],
        ["sabali nu", "else if"],
        ["sabali",    "else"],
    ],
    col_widths=[1.2, 4.1]
)
add_code_block(doc, """\
nu (x > 10) {
    Ibaga("big");
}
sabali nu (x == 10) {
    Ibaga("ten");
}
sabali {
    Ibaga("small");
}""")
add_body(doc, 'Conditions must be Pudno. No ; after closing }. Blocks create local scope.')

# ============================================================
# SCOPING
# ============================================================
add_heading_styled(doc, "SCOPING")
add_code_block(doc, """\
Bilang a dutokan-> 10;             // outer scope
nu (a > 5) {
    Bilang b dutokan-> 20;         // inner scope only
    Ibaga(a);                      // can access outer 'a'
}                                  // 'b' is NOT accessible here""")

# ============================================================
# ESCAPE SEQUENCES
# ============================================================
add_heading_styled(doc, "ESCAPE SEQUENCES")
add_body(doc, '\\n newline  |  \\t tab  |  \\\\ backslash  |  \\" quote')

# ============================================================
# FULL EXAMPLE
# ============================================================
add_heading_styled(doc, "COMPLETE EXAMPLE")
add_code_block(doc, """\
Sarsarita name dutokan-> Ikabil("What is your name? ");
Bilang x dutokan-> 10, y dutokan-> 5;
Gudua result dutokan-> ((x + y) * 2) / 3;
Bilang quotient dutokan-> y // 2;
Bilang remainder dutokan-> y % 2;
Ibaga("Hello, " + name);
Ibaga("Result: " + result);
nu (x > y) {
    Ibaga("x is greater");
} sabali {
    Ibaga("y is greater or equal");
}""")

# ============================================================
# HOW TO RUN
# ============================================================
add_heading_styled(doc, "HOW TO RUN")
add_code_block(doc, """\
python coderun.py file.elokano --out out.cpp --exe out.exe    // full pipeline
python elokano_analysis.py file.elokano                       // analysis only
python elokano_codegen.py file.elokano --out out.cpp           // generate C++ only""")
add_body(doc, 'Elokano compiles to C++ which is then compiled with g++/clang++/MSVC.', PEACH)

# ============================================================
# SAVE
# ============================================================
out_path = "Elokano_Cheatsheet.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
